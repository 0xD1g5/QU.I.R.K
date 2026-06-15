"""python-docx DOCX report renderer for QU.I.R.K. (Phase 100, FMT-03 / D-09..D-12).

Derives from the shared ExecContent + findings content model (D-10) — NOT an
HTML-to-DOCX conversion and NOT a parallel hand-built document.

python-docx is an optional extra ([docx]); when absent the function prints a
stderr advisory and returns False — mirrors the Playwright graceful-skip pattern
from render_pdf_report in html_renderer.py.

D-11: render_docx_report is called from write_reports on every report run.
D-12: structural fidelity (cover / exec / findings / roadmap / score sections,
       Word Heading 1/2 styles, native tables, editable default styles) and a
       clearly-marked logo placeholder paragraph for consultant editing.
"""
import logging
import os
import sys
from datetime import datetime, timezone
from typing import Any, List, Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Subscore label order — mirrors writer.py / html_renderer.py
# ---------------------------------------------------------------------------

_SUBSCORE_LABELS = [
    ("hygiene",         "Hygiene"),
    ("modern_tls",      "Modern TLS"),
    ("identity_trust",  "Identity"),
    ("agility_signals", "Agility"),
    ("data_at_rest",    "Data at Rest"),
    ("data_in_motion",  "Data in Motion"),
]


def _set_table_style(tbl, style_name: str = "Table Grid") -> None:
    """Assign a named style to a python-docx table, ignoring unknown-style errors.

    Research note A3: 'Table Grid' is a standard Word style but may be absent in
    some minimal document contexts. On any style-assignment error, leave the
    default table style in place — never crash on a missing style.
    """
    try:
        tbl.style = style_name
    except Exception:
        logger.warning(
            "Failed to apply table style %r — using default table style.",
            style_name,
            exc_info=True,
        )


def _set_col_widths(tbl, inches_list) -> None:
    """Pin explicit column widths (inches) so short headers like "Severity" / "Port"
    are not squeezed by long text columns and broken mid-word (FMT-02 — no truncation).

    python-docx requires width to be set on every cell in a column and autofit
    disabled for the widths to hold across Word and Google Docs. Best-effort:
    any failure (missing docx.shared, etc.) leaves the default sizing in place
    and never crashes the render (D-11 graceful contract).
    """
    try:
        from docx.shared import Inches
        tbl.autofit = False
        tbl.allow_autofit = False
        for row in tbl.rows:
            for idx, w in enumerate(inches_list):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(w)
    except Exception:
        logger.warning(
            "Failed to set column widths on table — using default sizing.",
            exc_info=True,
        )


def render_docx_report(
    path: str,
    cfg: Any,
    findings: List[dict],
    exec_content: "Any | None" = None,
) -> bool:
    """Write a structural Word DOCX report to *path*.

    Returns True on success, False if python-docx is not installed.

    Parameters
    ----------
    path : str
        Destination file path (e.g. /output/report-20260524-120000.docx).
    cfg : Any
        Config object with cfg.assessment.{name, report_owner, data_classification}.
    findings : list[dict]
        Finding dicts (keys: severity, title, host, port, description,
        recommendation, quantum_risk).
    exec_content : ExecContent | None
        Shared content model built by build_exec_content(). When None, falls back
        to empty narrative/roadmap/score sections (D-10 single pipeline).
    """
    # T-100-DEP: lazy import — MUST stay inside the function body.
    # Never import docx at module level (optional-extra import trap).
    try:
        from docx import Document
    except ImportError:
        print(
            "DOCX export skipped: python-docx is not installed. "
            "Install with: pip install quirk-scanner[docx]",
            file=sys.stderr,
        )
        return False

    # ---------------------------------------------------------------------------
    # D-10: route content through exec_content when present; fall back gracefully.
    # ---------------------------------------------------------------------------
    if exec_content is not None:
        narrative_lead = exec_content.narrative_lead
        narrative_drivers = exec_content.narrative_drivers or []
        top_risks = exec_content.top_risks or []
        roadmap_now = [r for r in exec_content.roadmap_items if r.phase == "NOW"]
        roadmap_next = [r for r in exec_content.roadmap_items if r.phase == "NEXT"]
        roadmap_later = [r for r in exec_content.roadmap_items if r.phase == "LATER"]
        subscores = exec_content.subscores or {}
        score_total = getattr(exec_content, "score_total", 0)
        raw_sum = getattr(exec_content, "raw_sum", 0)
    else:
        narrative_lead = None
        narrative_drivers = []
        top_risks = []
        roadmap_now = roadmap_next = roadmap_later = []
        subscores = {}
        score_total = 0
        raw_sum = 0

    # Phase 128 D-10: hardware advisory devices list — advisory-only, not scored
    hardware_devices = getattr(exec_content, "hardware_devices", []) if exec_content else []

    # ---------------------------------------------------------------------------
    # cfg access — double-getattr pattern (mirrors html_renderer.py lines 236-238)
    # ---------------------------------------------------------------------------
    org_name = getattr(getattr(cfg, "assessment", None), "name", "Unknown")
    report_owner = getattr(getattr(cfg, "assessment", None), "report_owner", "")
    data_classification = getattr(getattr(cfg, "assessment", None), "data_classification", "")
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # ---------------------------------------------------------------------------
    # Build Document
    # ---------------------------------------------------------------------------
    doc = Document()

    # FMT-02: the findings table is 7 columns wide — portrait Letter (~6.5" usable)
    # cannot give long headers ("Recommendation", "Severity") a full line once Word
    # cell margins are subtracted, forcing ugly mid-word wraps ("Recommendati on").
    # Landscape (~9.5" usable) gives every column room. Best-effort: never crash the
    # render if the section API is unavailable (D-11 graceful contract).
    try:
        from docx.enum.section import WD_ORIENT
        section = doc.sections[0]
        if section.page_width < section.page_height:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
    except Exception:
        logger.warning(
            "Failed to set landscape orientation — using default page orientation.",
            exc_info=True,
        )

    # ---- Cover block ----
    # 1. Logo placeholder paragraph (D-12 / 100-UI-SPEC.md §C — exact verbatim string)
    doc.add_paragraph("[ Insert organization logo here ]", style="Normal")

    # 2. Report title — Heading 1
    doc.add_heading("QU.I.R.K. Cryptographic Readiness Report", level=1)

    # 3. Organization sub-heading — Heading 2
    doc.add_heading(org_name, level=2)

    # 4. Metadata line — Normal
    doc.add_paragraph(
        f"Report Owner: {report_owner}  |  Date: {generated_at}"
        f"  |  Classification: {data_classification}",
        style="Normal",
    )

    # ---- Executive Summary section ----
    doc.add_heading("Executive Summary", level=1)

    if narrative_lead:
        doc.add_paragraph(narrative_lead, style="Normal")

    # Readiness Assessment sub-section
    doc.add_heading("Readiness Assessment", level=2)
    if narrative_drivers:
        for driver in narrative_drivers:
            doc.add_paragraph(str(driver), style="Normal")
    else:
        doc.add_paragraph("No readiness assessment drivers available.", style="Normal")

    # Score Decomposition sub-section (3-col table)
    doc.add_heading("Score Decomposition", level=2)
    score_decomp_tbl = doc.add_table(rows=1, cols=3)
    _set_table_style(score_decomp_tbl)
    hdr = score_decomp_tbl.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Score"
    hdr[2].text = "Budget"
    for key, label in _SUBSCORE_LABELS:
        row_cells = score_decomp_tbl.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(subscores.get(key, "—"))
        row_cells[2].text = "/25"

    # Priority Business Risks sub-section (2-col table)
    doc.add_heading("Priority Business Risks", level=2)
    risks_tbl = doc.add_table(rows=1, cols=2)
    _set_table_style(risks_tbl)
    risk_hdr = risks_tbl.rows[0].cells
    risk_hdr[0].text = "Risk"
    risk_hdr[1].text = "Business Impact"
    if top_risks:
        for risk in top_risks:
            row_cells = risks_tbl.add_row().cells
            row_cells[0].text = str(risk.risk_label)
            row_cells[1].text = str(risk.impact_sentence)
    else:
        row_cells = risks_tbl.add_row().cells
        row_cells[0].text = "No high-priority risks identified."
        row_cells[1].text = ""

    # Cross-surface parity (D-10 / EXEC-04): exclude advisory coverage_gap rows from
    # the findings tables, exactly as report.html.j2 does (rejectattr coverage_gap),
    # so a CISO-facing DOCX never shows scanner-advisory rows like "X scanner not
    # installed". Top Findings is capped at 10 to match the HTML [:10] limit.
    report_findings = [f for f in findings if f.get("category") != "coverage_gap"]

    # Top Findings sub-section (4-col table: Severity | Title | Host | Description)
    doc.add_heading("Top Findings", level=2)
    top_findings_tbl = doc.add_table(rows=1, cols=4)
    _set_table_style(top_findings_tbl)
    tf_hdr = top_findings_tbl.rows[0].cells
    tf_hdr[0].text = "Severity"
    tf_hdr[1].text = "Title"
    tf_hdr[2].text = "Host"
    tf_hdr[3].text = "Description"
    if report_findings:
        for f in report_findings[:10]:  # top 10 (parity with HTML)
            row_cells = top_findings_tbl.add_row().cells
            row_cells[0].text = str(f.get("severity", ""))
            row_cells[1].text = str(f.get("title", ""))
            row_cells[2].text = str(f.get("host", ""))
            row_cells[3].text = str(f.get("description", ""))
    else:
        row_cells = top_findings_tbl.add_row().cells
        row_cells[0].text = "No findings recorded for this scan."
        row_cells[1].text = ""
        row_cells[2].text = ""
        row_cells[3].text = ""

    # Pin widths so Severity/Host headers are not broken mid-word (FMT-02). 4 cols,
    # landscape Letter usable width ~9.5".
    _set_col_widths(top_findings_tbl, [0.9, 2.6, 1.4, 4.6])

    # ---- Findings section — 7-col table (100-UI-SPEC.md Word Table Column Contracts) ----
    doc.add_heading("Findings", level=1)
    findings_tbl = doc.add_table(rows=1, cols=7)
    _set_table_style(findings_tbl)
    f_hdr = findings_tbl.rows[0].cells
    for i, col_name in enumerate([
        "Severity", "Title", "Host", "Port",
        "Description", "Recommendation", "Quantum Risk",
    ]):
        f_hdr[i].text = col_name

    if report_findings:
        for f in report_findings:
            row_cells = findings_tbl.add_row().cells
            row_cells[0].text = str(f.get("severity", ""))
            row_cells[1].text = str(f.get("title", ""))
            row_cells[2].text = str(f.get("host", ""))
            row_cells[3].text = str(f.get("port", ""))
            row_cells[4].text = str(f.get("description", ""))
            row_cells[5].text = str(f.get("recommendation", ""))
            row_cells[6].text = str(f.get("quantum_risk", ""))
    else:
        # D-12 / 100-UI-SPEC.md §D: always write header + single empty-state data row
        row_cells = findings_tbl.add_row().cells
        row_cells[0].text = "No findings recorded for this scan."
        for i in range(1, 7):
            row_cells[i].text = ""

    # Pin widths: every header (incl. "Recommendation"/"Severity"/"Port") stays on
    # one line (FMT-02). Sums to ~9.5" (landscape Letter usable); long text columns
    # (Description/Recommendation/Quantum Risk) wrap their data internally.
    _set_col_widths(findings_tbl, [0.8, 1.5, 1.0, 0.6, 2.3, 1.6, 1.7])

    # ---- Remediation Roadmap section ----
    doc.add_heading("Remediation Roadmap", level=1)
    for phase_label, phase_items in [("NOW", roadmap_now), ("NEXT", roadmap_next), ("LATER", roadmap_later)]:
        doc.add_heading(phase_label, level=2)
        roadmap_tbl = doc.add_table(rows=1, cols=4)
        _set_table_style(roadmap_tbl)
        rm_hdr = roadmap_tbl.rows[0].cells
        rm_hdr[0].text = "Phase"
        rm_hdr[1].text = "Action"
        rm_hdr[2].text = "Rationale"
        rm_hdr[3].text = "Effort / Impact"
        if phase_items:
            for item in phase_items:
                row_cells = roadmap_tbl.add_row().cells
                row_cells[0].text = str(getattr(item, "phase", phase_label))
                row_cells[1].text = str(getattr(item, "title", ""))
                row_cells[2].text = str(getattr(item, "why", ""))
                effort = getattr(item, "effort", "")
                impact = getattr(item, "impact", "")
                row_cells[3].text = f"{effort} / {impact}" if (effort or impact) else ""
        else:
            row_cells = roadmap_tbl.add_row().cells
            row_cells[0].text = phase_label
            row_cells[1].text = f"No {phase_label} actions."
            row_cells[2].text = ""
            row_cells[3].text = ""

    # ---- Score Breakdown section ----
    doc.add_heading("Score Breakdown", level=1)
    # Rollup formula sentence
    doc.add_paragraph(
        f"{raw_sum} ÷ 1.5 = {score_total} / 100",
        style="Normal",
    )
    # Score Decomposition table (same as executive summary decomp)
    score_breakdown_tbl = doc.add_table(rows=1, cols=3)
    _set_table_style(score_breakdown_tbl)
    sb_hdr = score_breakdown_tbl.rows[0].cells
    sb_hdr[0].text = "Category"
    sb_hdr[1].text = "Score"
    sb_hdr[2].text = "Budget"
    for key, label in _SUBSCORE_LABELS:
        row_cells = score_breakdown_tbl.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(subscores.get(key, "—"))
        row_cells[2].text = "/25"

    # ---- Hardware PQC Advisory section (Phase 128 D-10) ----
    # Advisory-only — hardware findings do not affect the readiness score.
    # 7 columns: Tier | Vendor | Model | Host:Port | PQC Status | Confidence | EOL Date
    if hardware_devices:
        doc.add_heading("Hardware PQC Advisory (Not Scored)", level=2)
        doc.add_paragraph(
            "Advisory only — hardware findings do not affect the readiness score."
            " Listed for CNSA 2.0 migration planning purposes only.",
            style="Normal",
        )
        hw_tbl = doc.add_table(rows=1, cols=7)
        _set_table_style(hw_tbl)
        hw_hdr = hw_tbl.rows[0].cells
        for _i, _h in enumerate(
            ["Tier", "Vendor", "Model", "Host:Port", "PQC Status", "Confidence", "EOL Date"]
        ):
            hw_hdr[_i].text = _h
        for _d in hardware_devices:
            _row = hw_tbl.add_row().cells
            _row[0].text = _d.get("remediation_tier", "")
            _row[1].text = _d.get("vendor", "")
            _row[2].text = _d.get("model") or "Unknown"
            _row[3].text = f"{_d.get('host', '')}:{_d.get('port', '')}"
            _row[4].text = _d.get("pqc_status", "")
            _row[5].text = _d.get("confidence", "")
            _row[6].text = _d.get("eol_date") or "—"
        _set_col_widths(hw_tbl, [0.9, 1.4, 1.4, 1.5, 1.1, 1.1, 0.9])

    # ---------------------------------------------------------------------------
    # Save document
    # ---------------------------------------------------------------------------
    os.makedirs(os.path.dirname(path) if os.path.dirname(path) else ".", exist_ok=True)
    try:
        doc.save(path)
    except Exception as e:
        logger.warning(
            "DOCX export failed while writing to output path.",
            exc_info=True,
        )
        print(
            f"DOCX export failed while writing {path}: {e}",
            file=sys.stderr,
        )
        return False
    print(f"DOCX report written to {path}", file=sys.stderr)
    return True
