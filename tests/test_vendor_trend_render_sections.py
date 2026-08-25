"""Phase 161 Plan 05 (HWLC-19) — "Vendor PQC Status Trends" HTML + DOCX sections.

Plan 161-02 delivered the CLI surface and the data source; this covers the other
two export surfaces required by ROADMAP success criterion 4.

Vendor-trend rows are catalog-level: vendor / event_type / old_value / new_value
/ detected_at / confirmed_at, with no host, port, direction or severity. The
table is therefore four columns, not the drift table's five.

Mirrors tests/test_docx_renderer_drift.py's Document-open + text-readback idiom.
"""
from __future__ import annotations

from types import SimpleNamespace

import pytest


def _trend(**overrides) -> dict:
    ev = {
        "vendor": "Cisco",
        "event_type": "pqc_status_change",
        "old_value": "0/3 models PQC-capable",
        "new_value": "1/3 models PQC-capable",
        "detected_at": "2026-08-20",
        "confirmed_at": "2026-08-20",
    }
    ev.update(overrides)
    return ev


def _make_minimal_cfg(tmpdir="/tmp/quirk_test_vendor_trend"):
    return SimpleNamespace(
        assessment=SimpleNamespace(
            name="Vendor Trend Test Org",
            report_owner="Trend Tester",
            data_classification="CONFIDENTIAL",
            timezone="UTC",
            logo_path=None,
        ),
        output=SimpleNamespace(directory=tmpdir),
    )


def _exec_content(vendor_pqc_trends):
    from quirk.reports.content_model import ExecContent

    return ExecContent(
        narrative_lead="Test narrative lead.",
        narrative_drivers=[],
        top_risks=[],
        roadmap_items=[],
        score_total=70,
        score_band="FAIR",
        subscores={},
        raw_sum=0,
        sev_counts={},
        vendor_pqc_trends=vendor_pqc_trends,
    )


# ---------------------------------------------------------------------------
# Task 1 — render_vendor_trend_section()
# ---------------------------------------------------------------------------


class TestRenderVendorTrendSection:
    def test_empty_list_returns_empty_string(self):
        from quirk.reports.html_renderer import render_vendor_trend_section

        assert render_vendor_trend_section([]) == ""

    def test_none_returns_empty_string(self):
        """No orphan heading when the loader produced nothing."""
        from quirk.reports.html_renderer import render_vendor_trend_section

        assert render_vendor_trend_section(None) == ""

    def test_populated_section_has_heading_caption_and_four_columns(self):
        from quirk.reports.html_renderer import (
            VENDOR_TREND_ADVISORY_CAPTION,
            render_vendor_trend_section,
        )

        html = render_vendor_trend_section([_trend()])
        assert "Vendor PQC Status Trends" in html
        assert VENDOR_TREND_ADVISORY_CAPTION in html
        for label in ("Vendor", "Change", "Transition", "Detected"):
            assert f">{label}</th>" in html, f"missing th cell {label}"

    def test_no_device_or_direction_column(self):
        """Vendor-scoped data has no host/port and no direction (D-02)."""
        from quirk.reports.html_renderer import render_vendor_trend_section

        html = render_vendor_trend_section([_trend()])
        assert ">Direction</th>" not in html
        assert ">Device</th>" not in html

    def test_hostile_vendor_value_is_escaped(self):
        """T-161-18: the template interpolates with | safe, so escaping here is the defence."""
        from quirk.reports.html_renderer import render_vendor_trend_section

        html = render_vendor_trend_section([_trend(vendor="<script>x</script>")])
        assert "<script>" not in html, "unescaped markup reached the report"
        assert "&lt;script&gt;" in html

    def test_unknown_event_type_falls_through_to_raw_value(self):
        from quirk.reports.html_renderer import render_vendor_trend_section

        html = render_vendor_trend_section([_trend(event_type="brand_new_type")])
        assert "brand_new_type" in html

    def test_null_transition_side_renders_em_dash_not_none(self):
        from quirk.reports.html_renderer import render_vendor_trend_section

        html = render_vendor_trend_section([_trend(old_value=None)])
        assert "—" in html
        assert ">None" not in html
        assert "None &#x2192;" not in html


# ---------------------------------------------------------------------------
# Task 2 — end-to-end through render_html_report()
# ---------------------------------------------------------------------------


class TestHtmlReportWiring:
    def test_section_present_when_trends_exist(self, tmp_path):
        from quirk.reports.html_renderer import render_html_report

        path = str(tmp_path / "with_trends.html")
        render_html_report(
            path=path,
            cfg=_make_minimal_cfg(str(tmp_path)),
            endpoints=[],
            findings=[],
            score={"total": 70, "subscores": {}, "drivers": []},
            conf={"confidence": 80, "confidence_factors": {}},
            roadmap_items=[],
            exec_content=_exec_content([_trend()]),
        )
        html = open(path, encoding="utf-8").read()
        assert "Vendor PQC Status Trends" in html

    def test_section_absent_when_no_trends(self, tmp_path):
        from quirk.reports.html_renderer import render_html_report

        path = str(tmp_path / "without_trends.html")
        render_html_report(
            path=path,
            cfg=_make_minimal_cfg(str(tmp_path)),
            endpoints=[],
            findings=[],
            score={"total": 70, "subscores": {}, "drivers": []},
            conf={"confidence": 80, "confidence_factors": {}},
            roadmap_items=[],
            exec_content=_exec_content([]),
        )
        html = open(path, encoding="utf-8").read()
        assert "Vendor PQC Status Trends" not in html


# ---------------------------------------------------------------------------
# Task 3 — DOCX section
# ---------------------------------------------------------------------------


def _all_paragraph_texts(doc):
    return [p.text for p in doc.paragraphs]


class TestDocxVendorTrendSection:
    def test_heading_present(self, tmp_path):
        from docx import Document
        from quirk.reports.docx_renderer import render_docx_report

        path = str(tmp_path / "trend_heading.docx")
        render_docx_report(
            path=path, cfg=_make_minimal_cfg(str(tmp_path)), findings=[],
            exec_content=_exec_content([_trend()]),
        )
        headings = [
            p.text for p in Document(path).paragraphs
            if p.style.name.startswith("Heading")
        ]
        assert "Vendor PQC Status Trends" in headings

    def test_caption_matches_the_html_constant_exactly(self, tmp_path):
        from docx import Document
        from quirk.reports.docx_renderer import render_docx_report
        from quirk.reports.html_renderer import VENDOR_TREND_ADVISORY_CAPTION

        path = str(tmp_path / "trend_caption.docx")
        render_docx_report(
            path=path, cfg=_make_minimal_cfg(str(tmp_path)), findings=[],
            exec_content=_exec_content([_trend()]),
        )
        assert VENDOR_TREND_ADVISORY_CAPTION in _all_paragraph_texts(Document(path))

    def test_table_has_four_columns_with_the_right_headers(self, tmp_path):
        from docx import Document
        from quirk.reports.docx_renderer import render_docx_report

        path = str(tmp_path / "trend_table.docx")
        render_docx_report(
            path=path, cfg=_make_minimal_cfg(str(tmp_path)), findings=[],
            exec_content=_exec_content([_trend()]),
        )
        tables = Document(path).tables
        match = [
            t for t in tables
            if [c.text for c in t.rows[0].cells]
            == ["Vendor", "Change", "Transition", "Detected"]
        ]
        assert len(match) == 1, "vendor-trend table not found with its 4 headers"
        assert len(match[0].columns) == 4

    def test_zero_trends_emits_no_heading_and_no_empty_table(self, tmp_path):
        from docx import Document
        from quirk.reports.docx_renderer import render_docx_report

        path = str(tmp_path / "trend_none.docx")
        render_docx_report(
            path=path, cfg=_make_minimal_cfg(str(tmp_path)), findings=[],
            exec_content=_exec_content([]),
        )
        doc = Document(path)
        assert "Vendor PQC Status Trends" not in _all_paragraph_texts(doc)
        assert not [
            t for t in doc.tables
            if [c.text for c in t.rows[0].cells]
            == ["Vendor", "Change", "Transition", "Detected"]
        ]

    def test_exec_content_none_does_not_raise(self, tmp_path):
        from quirk.reports.docx_renderer import render_docx_report

        path = str(tmp_path / "trend_no_exec.docx")
        assert render_docx_report(
            path=path, cfg=_make_minimal_cfg(str(tmp_path)), findings=[],
            exec_content=None,
        ) is True


# ---------------------------------------------------------------------------
# Cross-surface caption parity — fails loudly if the three surfaces drift apart
# ---------------------------------------------------------------------------


def test_advisory_caption_is_identical_across_all_three_surfaces():
    import inspect

    from quirk.reports import technical
    from quirk.reports.docx_renderer import _VENDOR_TREND_ADVISORY_CAPTION
    from quirk.reports.html_renderer import VENDOR_TREND_ADVISORY_CAPTION

    assert VENDOR_TREND_ADVISORY_CAPTION == _VENDOR_TREND_ADVISORY_CAPTION, (
        "HWLC-19: the HTML and DOCX advisory captions have drifted apart"
    )
    assert VENDOR_TREND_ADVISORY_CAPTION in inspect.getsource(technical), (
        "HWLC-19: the CLI report's caption no longer matches HTML/DOCX"
    )
