"""Phase 181 Plan 02 (SURF-02 / T-181-04, T-181-06, T-181-09, T-181-10, T-181-11) —
remediation burndown rendering across the CLI markdown, HTML, and DOCX report surfaces.

Every honest state Phases 178-180 built — `not_observed`, `resurfaced`, `unmapped`,
refusal — must survive rendering unchanged. A surface that silently drops one of them
re-creates exactly the overclaiming this milestone exists to prevent.

CAPTION MECHANISM (read this before "DRY-ing up" anything below): the byte-identical
advisory caption across CLI/HTML/DOCX is enforced by THREE INDEPENDENTLY DUPLICATED
per-renderer string constants held equal by a test — NOT by a shared constant in
content_model.py. This mirrors the real Phase 161 HWLC-19 precedent
(`tests/test_vendor_trend_render_sections.py::test_advisory_caption_is_identical_across_all_three_surfaces`,
`quirk/reports/html_renderer.py:503-508`, `quirk/reports/docx_renderer.py:114-121`,
`quirk/reports/technical.py:37`). A future reader who "simplifies" this into one shared
constant in content_model.py deletes the gate this file exists to hold open — an
initial 181-CONTEXT.md draft made exactly this mistake and its own addendum retracted
it. Do not repeat it.

REFUSAL STATEMENT is the deliberate asymmetry: it is NOT a fourth duplicated constant.
It is computed ONCE in writer.py from the comparability-refusal axis and printed
verbatim, unmodified, by all three renderers. `grep -c "Closure not computed: "` over
this file's fixtures is expected to find the phrase originating from a single
computation, never three separately worded copies.

This plan writes TESTS ONLY, before any of the interfaces below exist. Plans 181-05
(ExecContent.burndown / .closure_refusal payload) and 181-06 (the three renderers)
implement against this file. It must collect cleanly and fail RED today.
"""
from __future__ import annotations

import re
from types import SimpleNamespace

import pytest

# ---------------------------------------------------------------------------
# Shared fixtures
# ---------------------------------------------------------------------------


def _burndown(**overrides) -> dict:
    """Three-bucket shape verbatim from `compute_burndown()` (quirk/intelligence/burndown.py).

    `unmapped` deliberately carries a NON-ZERO `open` count so a renderer that skips
    zero-count buckets does not accidentally pass the visibility assertions below.
    """
    base = {
        "key_establishment": {
            "date": "2030-12-31",
            "standard": "FIPS 203 (ML-KEM)",
            "fingerprints": 10,
            "open": 3,
            "closed": 5,
            "not_observed": 1,
            "resurfaced": 1,
            "open_like": 4,
        },
        "digital_signature": {
            "date": "2031-12-31",
            "standard": "FIPS 186-5 (DSS)",
            "fingerprints": 8,
            "open": 2,
            "closed": 4,
            "not_observed": 1,
            "resurfaced": 1,
            "open_like": 3,
        },
        "unmapped": {
            "date": None,
            "standard": None,
            "fingerprints": 6,
            "open": 6,
            "closed": 0,
            "not_observed": 0,
            "resurfaced": 0,
            "open_like": 6,
        },
    }
    base.update(overrides)
    return base


def _refusal(reason_key: str, axis: str) -> dict:
    """Build a closure_refusal dict matching the locked shape in 181-CONTEXT.md's
    interfaces block. `statement` mirrors writer.py's single-computation contract —
    "Closure not computed: {axis}." — so this fixture does not itself duplicate the
    wording renderers must reproduce verbatim.
    """
    return {
        "refused": True,
        "reason_key": reason_key,
        "axis": axis,
        "statement": f"Closure not computed: {axis}.",
    }


# The five scan-level refusal reasons from `scans_are_comparable`'s fixed comparability
# ladder (quirk/intelligence/closure.py COMPARABILITY_REASONS, minus "comparable"),
# reason-keyed to match `_COUNTER_KEYS`'s "refused_*" naming convention. Axis phrasing
# matches docs/operators-guide.md §"Why an item reads not_observed" vocabulary.
REFUSAL_REASONS = [
    ("refused_no_prior", "no comparable prior scan exists"),
    ("refused_missing_signature", "scope signature is missing for one of the two scans"),
    ("refused_signature_version_gap", "scope signature versions differ between the two scans"),
    (
        "refused_missing_target_set_digest",
        "target set digest is missing for one of the two scans",
    ),
    ("refused_scope_mismatch", "scan scope differs from the prior scan"),
]


def _make_minimal_cfg(tmpdir="/tmp/quirk_test_burndown"):
    return SimpleNamespace(
        assessment=SimpleNamespace(
            name="Burndown Test Org",
            report_owner="Burndown Tester",
            data_classification="CONFIDENTIAL",
            timezone="UTC",
            logo_path=None,
        ),
        output=SimpleNamespace(directory=tmpdir),
    )


def _exec_content(burndown=None, closure_refusal=None):
    from quirk.reports.content_model import ExecContent

    return ExecContent(
        narrative_lead="Test narrative lead.",
        narrative_drivers=[],
        top_risks=[],
        roadmap_items=[],
        score_total=70,
        score_band="FAIR",
        subscores={},
        raw_sum=0,
        sev_counts={},
        burndown=burndown if burndown is not None else {},
        closure_refusal=closure_refusal if closure_refusal is not None else {},
    )


def _all_paragraph_texts(doc):
    return [p.text for p in doc.paragraphs]


def _all_table_texts(doc):
    cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cells.append(cell.text)
    return cells


def _docx_full_text(doc) -> str:
    return "\n".join(_all_paragraph_texts(doc) + _all_table_texts(doc))


_FORBIDDEN_AGGREGATE_TOKENS = ("total", "overall", "aggregate", "remediated")


def _assert_no_forbidden_aggregate_tokens(text: str, *, surface: str) -> None:
    lowered = text.lower()
    for token in _FORBIDDEN_AGGREGATE_TOKENS:
        assert token not in lowered, (
            f"{surface}: forbidden aggregate/percentage token {token!r} found in "
            "burndown output — CLOSE-03 eliminated the single scalar; buckets overlap "
            "by design so any sum double-counts"
        )
    assert "%" not in text, f"{surface}: burndown output must never render a percentage"


# ---------------------------------------------------------------------------
# Task 1 — caption parity across all three surfaces + no-severity payload contract
# ---------------------------------------------------------------------------


class TestExecContentBurndownFields:
    def test_exec_content_defaults_are_empty(self):
        """Every pre-existing ExecContent(...) construction in the suite keeps working."""
        content = _exec_content()
        assert content.burndown == {}
        assert content.closure_refusal == {}

    def test_burndown_payload_declares_no_severity_host_or_port(self):
        """T-181-04: the absent `severity` key is the STRUCTURAL mechanism that keeps
        this payload off the findings chokepoint (_build_finding() / findings_evaluator.py)
        that feeds scoring — it is not decoration. The absent `host`/`port` keys keep
        estate identity out of the published deliverable, matching the `vendor_pqc_trends`
        / `eol_forecast` contract in content_model.py:100-112.
        """
        import inspect

        from quirk.reports import content_model

        source_lines = inspect.getsource(content_model).splitlines()
        # Grep-gate hygiene: drop comment-only lines first — the field's own explanatory
        # comment legitimately contains the word "severity" (see the docstring above).
        code_lines = [ln for ln in source_lines if not ln.strip().startswith("#")]
        code_source = "\n".join(code_lines)

        for field_name in ("burndown", "closure_refusal"):
            match = re.search(rf"^\s*{field_name}\s*:.*$", code_source, re.MULTILINE)
            assert match is not None, (
                f"content_model.py has no `{field_name}:` field declaration yet "
                "(expected pre-181-05 — this assertion is the RED signal)"
            )
            declaration = match.group(0)
            for forbidden in ("severity", "host", "port"):
                assert forbidden not in declaration.lower(), (
                    f"`{field_name}` field declaration must not carry a `{forbidden}` "
                    "key — that would route this advisory payload into the findings "
                    "chokepoint that feeds scoring"
                )


def test_advisory_caption_is_identical_across_all_three_surfaces():
    """Mirrors test_vendor_trend_render_sections.py's node of the same name exactly.
    SURF-02: three surfaces with three independently-worded advisory captions lets a
    client be shown a weaker caveat in whichever format they happen to read.
    """
    import inspect

    from quirk.reports import technical
    from quirk.reports.docx_renderer import _BURNDOWN_ADVISORY_CAPTION
    from quirk.reports.html_renderer import BURNDOWN_ADVISORY_CAPTION

    assert BURNDOWN_ADVISORY_CAPTION == _BURNDOWN_ADVISORY_CAPTION, (
        "SURF-02: the HTML and DOCX burndown advisory captions have drifted apart"
    )
    assert BURNDOWN_ADVISORY_CAPTION in inspect.getsource(technical), (
        "SURF-02: the CLI report's burndown caption no longer matches HTML/DOCX"
    )
    assert "Advisory" in BURNDOWN_ADVISORY_CAPTION


# ---------------------------------------------------------------------------
# Task 2 — per-deadline sections, unmapped visible, no aggregate anywhere
# ---------------------------------------------------------------------------


class TestPerDeadlineSections:
    def test_html_section_shows_both_deadlines_and_unmapped(self):
        from quirk.reports.html_renderer import render_burndown_section

        html = render_burndown_section(_burndown(), None)
        assert "Remediation Burndown" in html
        assert "2030-12-31" in html
        assert "2031-12-31" in html
        assert "unmapped" in html.lower()
        for state_label in ("open", "closed", "not_observed", "resurfaced"):
            assert state_label in html.lower() or state_label.replace("_", " ") in html.lower(), (
                f"HTML burndown section missing state label {state_label!r}"
            )

    def test_cli_markdown_shows_both_deadlines_and_unmapped(self):
        from quirk.reports.technical import build_tech_markdown

        md = build_tech_markdown(
            _make_minimal_cfg(), [], [], burndown=_burndown(), closure_refusal=None
        )
        assert "2030-12-31" in md
        assert "2031-12-31" in md
        assert "unmapped" in md.lower()

    def test_docx_shows_both_deadlines_and_unmapped(self, tmp_path):
        from docx import Document
        from quirk.reports.docx_renderer import render_docx_report

        path = str(tmp_path / "burndown_deadlines.docx")
        render_docx_report(
            path=path,
            cfg=_make_minimal_cfg(str(tmp_path)),
            findings=[],
            exec_content=_exec_content(burndown=_burndown()),
        )
        text = _docx_full_text(Document(path))
        assert "2030-12-31" in text
        assert "2031-12-31" in text
        assert "unmapped" in text.lower()


def test_no_aggregate_or_percentage_on_any_surface(tmp_path):
    """CLOSE-03: BURNDOWN_BUCKETS overlap by design and are never summed. A headline
    total or "% remediated" over deliberately overlapping buckets is a double-counted
    number that would be quoted back to a client as fact — this reintroduces exactly
    the single scalar CLOSE-03 eliminated.
    """
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report
    from quirk.reports.html_renderer import render_burndown_section
    from quirk.reports.technical import build_tech_markdown

    html = render_burndown_section(_burndown(), None)
    _assert_no_forbidden_aggregate_tokens(html, surface="HTML")

    md = build_tech_markdown(
        _make_minimal_cfg(), [], [], burndown=_burndown(), closure_refusal=None
    )
    # Scope to the burndown section only — surrounding report boilerplate (score
    # narrative, etc.) is out of scope for this assertion.
    heading_idx = md.find("Remediation Burndown")
    assert heading_idx != -1, "CLI markdown missing the Remediation Burndown heading"
    burndown_block = md[heading_idx:]
    next_heading = re.search(r"\n#{1,2} ", burndown_block[1:])
    if next_heading:
        burndown_block = burndown_block[: next_heading.start() + 1]
    _assert_no_forbidden_aggregate_tokens(burndown_block, surface="CLI markdown")

    path = str(tmp_path / "burndown_no_aggregate.docx")
    render_docx_report(
        path=path,
        cfg=_make_minimal_cfg(str(tmp_path)),
        findings=[],
        exec_content=_exec_content(burndown=_burndown()),
    )
    doc = Document(path)
    heading_seen = False
    docx_block_parts = []
    for p in doc.paragraphs:
        if p.text == "Remediation Burndown":
            heading_seen = True
            continue
        if heading_seen:
            if p.style.name.startswith("Heading"):
                break
            docx_block_parts.append(p.text)
    docx_block_parts.extend(_all_table_texts(doc))
    _assert_no_forbidden_aggregate_tokens("\n".join(docx_block_parts), surface="DOCX")


class TestEmptyBurndownRendersNothing:
    def test_html_empty_burndown_returns_empty_string(self):
        from quirk.reports.html_renderer import render_burndown_section

        assert render_burndown_section({}, None) == ""

    def test_docx_empty_burndown_emits_no_heading(self, tmp_path):
        from docx import Document
        from quirk.reports.docx_renderer import render_docx_report

        path = str(tmp_path / "burndown_empty.docx")
        render_docx_report(
            path=path,
            cfg=_make_minimal_cfg(str(tmp_path)),
            findings=[],
            exec_content=_exec_content(burndown={}),
        )
        headings = [
            p.text
            for p in Document(path).paragraphs
            if p.style.name.startswith("Heading")
        ]
        assert "Remediation Burndown" not in headings


# ---------------------------------------------------------------------------
# Task 3 — a refused scan is stated explicitly, with the differing axis named
# ---------------------------------------------------------------------------


class TestRefusalDisclosure:
    def test_refusal_statement_is_rendered_on_every_surface(self, tmp_path):
        """A refused scan is NOT silence. All three surfaces render the SAME `statement`
        string, byte-identical — writer.py computes it once, renderers print it verbatim.
        """
        from docx import Document
        from quirk.reports.docx_renderer import render_docx_report
        from quirk.reports.html_renderer import render_burndown_section
        from quirk.reports.technical import build_tech_markdown

        refusal = _refusal("refused_scope_mismatch", "scan scope differs from the prior scan")
        statement = refusal["statement"]

        html = render_burndown_section({}, refusal)
        assert statement in html, "HTML did not render the refusal statement verbatim"

        md = build_tech_markdown(
            _make_minimal_cfg(), [], [], burndown={}, closure_refusal=refusal
        )
        assert statement in md, "CLI markdown did not render the refusal statement verbatim"

        path = str(tmp_path / "burndown_refusal.docx")
        render_docx_report(
            path=path,
            cfg=_make_minimal_cfg(str(tmp_path)),
            findings=[],
            exec_content=_exec_content(burndown={}, closure_refusal=refusal),
        )
        docx_text = _docx_full_text(Document(path))
        assert statement in docx_text, "DOCX did not render the refusal statement verbatim"

    def test_refusal_is_never_presented_as_zero_closed(self, tmp_path):
        """The false-negative twin of the false-closure risk: a burndown that reads
        "clean" when it means "unchecked" is as dangerous as one that reads "closed"
        when it means "never rechecked".
        """
        from docx import Document
        from quirk.reports.docx_renderer import render_docx_report
        from quirk.reports.html_renderer import render_burndown_section
        from quirk.reports.technical import build_tech_markdown

        refusal = _refusal("refused_scope_mismatch", "scan scope differs from the prior scan")
        forbidden = ("nothing closed", "0 closed", "no open items")

        html = render_burndown_section({}, refusal).lower()
        md = build_tech_markdown(
            _make_minimal_cfg(), [], [], burndown={}, closure_refusal=refusal
        ).lower()

        path = str(tmp_path / "burndown_refusal_not_clean.docx")
        render_docx_report(
            path=path,
            cfg=_make_minimal_cfg(str(tmp_path)),
            findings=[],
            exec_content=_exec_content(burndown={}, closure_refusal=refusal),
        )
        docx_text = _docx_full_text(Document(path)).lower()

        for surface_name, text in (("HTML", html), ("CLI markdown", md), ("DOCX", docx_text)):
            for phrase in forbidden:
                assert phrase not in text, (
                    f"{surface_name}: a refusal must never be presentable as a clean "
                    f"result — found forbidden phrase {phrase!r}"
                )

    @pytest.mark.parametrize("reason_key,axis", REFUSAL_REASONS)
    def test_refusal_names_the_differing_axis(self, reason_key, axis):
        """'Closure not computed' alone is insufficient — the client must be told WHICH
        dimension differed. Parametrized over all five scan-level refusal reasons so a
        future reason added to `_COUNTER_KEYS` without an axis phrase fails loudly
        rather than rendering an empty explanation.
        """
        from quirk.reports.html_renderer import render_burndown_section

        refusal = _refusal(reason_key, axis)
        html = render_burndown_section({}, refusal)
        assert axis in html, (
            f"burndown section for {reason_key!r} did not name the differing axis "
            f"{axis!r} — 'closure not computed' alone does not tell the client which "
            "dimension differed"
        )
