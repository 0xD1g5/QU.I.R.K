"""Phase 146 Plan 03 — undetermined-host disclosure render-parity tests (DISC-07 / D-08/D-09/D-10).

Presence-based tests (per feedback_report_render_tests_presence_not_appearance): assert the
"Hosts undetermined (unreachable/filtered)" label string and the undetermined count appear in
the CLI markdown, HTML, and DOCX output. These check content presence only — not visual
appearance, table structure, or rendered position — they only check that the shared
ExecContent.undetermined_hosts_count / .undetermined_hosts_breakdown fields are surfaced
verbatim on every renderer.

Node IDs:
  test_markdown_shows_undetermined_headline_and_breakdown
  test_markdown_zero_count_omits_breakdown
  test_html_shows_undetermined_headline_and_count
  test_docx_shows_undetermined_headline_and_count
  test_cross_surface_parity_undetermined_count
  test_compute_undetermined_hosts_excludes_non_discovery_rows
"""
from __future__ import annotations

import os
from types import SimpleNamespace

from quirk.reports.content_model import ExecContent

_HEADLINE = "Hosts undetermined (unreachable/filtered)"


def _make_minimal_cfg(tmp_dir: str = "/tmp/quirk_test_undetermined_hosts"):
    """Minimal cfg-like namespace for renderer calls (mirrors test_cross_surface_parity.py)."""
    return SimpleNamespace(
        assessment=SimpleNamespace(
            name="Undetermined Hosts Test Org",
            report_owner="Parity Tester",
            data_classification="CONFIDENTIAL",
            timezone="UTC",
            logo_path=None,
        ),
        output=SimpleNamespace(directory=tmp_dir),
        intelligence=SimpleNamespace(
            profile="balanced",
            calibration_overrides=None,
        ),
    )


def _make_exec_content(count: int, breakdown: dict) -> ExecContent:
    return ExecContent(
        narrative_lead="Test narrative lead for undetermined-host disclosure.",
        narrative_drivers=[],
        top_risks=[],
        roadmap_items=[],
        score_total=70,
        score_band="FAIR",
        subscores={},
        raw_sum=0,
        sev_counts={},
        undetermined_hosts_count=count,
        undetermined_hosts_breakdown=breakdown,
    )


# ---------------------------------------------------------------------------
# Markdown (CLI)
# ---------------------------------------------------------------------------


def test_markdown_shows_undetermined_headline_and_breakdown():
    """build_exec_markdown renders the headline + count and both breakdown sub-bullets."""
    from quirk.reports.executive import build_exec_markdown

    exec_content = _make_exec_content(3, {"discovery_exception": 1, "liveness_skip": 2})
    cfg = _make_minimal_cfg()

    output = build_exec_markdown(
        cfg=cfg,
        endpoints=[],
        findings=[],
        exec_content=exec_content,
    )

    assert _HEADLINE in output, f"headline missing from markdown output: {output[:600]!r}"
    assert "3" in output.split(_HEADLINE, 1)[1].split("\n", 1)[0], (
        "undetermined count not adjacent to headline in markdown output"
    )
    assert "no response to liveness pre-pass" in output
    assert "discovery batch errors" in output


def test_markdown_zero_count_omits_breakdown():
    """count == 0 still renders the headline (so 'zero' is distinguishable from 'not measured')
    but omits the breakdown sub-bullets."""
    from quirk.reports.executive import build_exec_markdown

    exec_content = _make_exec_content(0, {"discovery_exception": 0, "liveness_skip": 0})
    cfg = _make_minimal_cfg()

    output = build_exec_markdown(
        cfg=cfg,
        endpoints=[],
        findings=[],
        exec_content=exec_content,
    )

    assert _HEADLINE in output
    assert "no response to liveness pre-pass" not in output
    assert "discovery batch errors" not in output


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


def test_html_shows_undetermined_headline_and_count(tmp_path):
    """render_html_report writes a file whose text contains the label and the count."""
    from quirk.reports.html_renderer import render_html_report

    exec_content = _make_exec_content(2, {"discovery_exception": 2, "liveness_skip": 0})
    cfg = _make_minimal_cfg()
    html_path = os.path.join(str(tmp_path), "undetermined-test.html")

    render_html_report(
        path=html_path,
        cfg=cfg,
        endpoints=[],
        findings=[],
        score={"total": 70, "subscores": {}, "drivers": []},
        conf={"confidence": 60, "confidence_factors": {}},
        roadmap_items=[],
        exec_content=exec_content,
    )

    html_output = open(html_path, encoding="utf-8").read()
    assert "Hosts undetermined" in html_output
    assert "<td>2</td>" in html_output


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_docx_shows_undetermined_headline_and_count(tmp_path):
    """render_docx_report produces a document whose paragraph text contains the label + count."""
    import pytest

    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed — skipping DOCX check")

    from quirk.reports.docx_renderer import render_docx_report

    exec_content = _make_exec_content(5, {"discovery_exception": 1, "liveness_skip": 4})
    cfg = _make_minimal_cfg()
    docx_path = os.path.join(str(tmp_path), "undetermined-test.docx")

    result = render_docx_report(
        path=docx_path, cfg=cfg, findings=[], exec_content=exec_content
    )
    if result is False:
        pytest.skip("python-docx not installed — skipping DOCX check")

    doc = Document(docx_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert _HEADLINE in full_text
    assert "5" in full_text
    assert "no response to liveness pre-pass" in full_text
    assert "discovery batch errors" in full_text


# ---------------------------------------------------------------------------
# Cross-surface parity
# ---------------------------------------------------------------------------


def test_cross_surface_parity_undetermined_count(tmp_path):
    """The same count string appears in markdown, HTML, and DOCX outputs built from
    one ExecContent instance (D-08 shared-field guarantee)."""
    import pytest

    from quirk.reports.executive import build_exec_markdown
    from quirk.reports.html_renderer import render_html_report
    from quirk.reports.docx_renderer import render_docx_report

    exec_content = _make_exec_content(7, {"discovery_exception": 3, "liveness_skip": 4})
    cfg = _make_minimal_cfg()

    md_output = build_exec_markdown(
        cfg=cfg, endpoints=[], findings=[], exec_content=exec_content
    )
    assert _HEADLINE in md_output
    assert "7" in md_output

    html_path = os.path.join(str(tmp_path), "parity-undetermined.html")
    render_html_report(
        path=html_path,
        cfg=cfg,
        endpoints=[],
        findings=[],
        score={"total": 70, "subscores": {}, "drivers": []},
        conf={"confidence": 60, "confidence_factors": {}},
        roadmap_items=[],
        exec_content=exec_content,
    )
    html_output = open(html_path, encoding="utf-8").read()
    assert "Hosts undetermined" in html_output
    assert "<td>7</td>" in html_output

    docx_path = os.path.join(str(tmp_path), "parity-undetermined.docx")
    result = render_docx_report(
        path=docx_path, cfg=cfg, findings=[], exec_content=exec_content
    )
    if result is False:
        pytest.skip("python-docx not installed — skipping DOCX parity leg")

    from docx import Document

    doc = Document(docx_path)
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    assert _HEADLINE in docx_text
    assert "7" in docx_text


# ---------------------------------------------------------------------------
# Pitfall-3 regression: non-discovery rows excluded
# ---------------------------------------------------------------------------


def test_compute_undetermined_hosts_excludes_non_discovery_rows():
    """Locks the Pitfall-3 invariant: a live-host TLS/SSH/API scan error (port != 0) and a
    non-discovery error category (even at port 0) must NOT be counted as undetermined."""
    from quirk.reports.writer import _compute_undetermined_hosts

    endpoints = [
        SimpleNamespace(port=443, scan_error_category="discovery_exception"),
        SimpleNamespace(port=0, scan_error_category="missing_extra"),
    ]
    count, breakdown = _compute_undetermined_hosts(endpoints)
    assert count == 0
    assert breakdown == {"discovery_exception": 0, "liveness_skip": 0}


def test_compute_undetermined_hosts_excludes_generic_wrapped_phase_exception():
    """CR-01 regression: _wrapped_phase()'s generic "exception" category — used by every
    non-discovery scanner stage (TLS, SSH, JWT, container, ...) — must NOT be counted as
    undetermined, even at port=0, since its host field is a scanner label, not a target host."""
    from quirk.reports.writer import _compute_undetermined_hosts

    endpoints = [
        SimpleNamespace(host="tls_scanner", port=0, scan_error_category="exception"),
    ]
    count, breakdown = _compute_undetermined_hosts(endpoints)
    assert count == 0
    assert breakdown == {"discovery_exception": 0, "liveness_skip": 0}
