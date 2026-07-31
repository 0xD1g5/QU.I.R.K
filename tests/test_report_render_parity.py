"""Phase 139 Plan 05 — SNMP column render-parity tests (SNMPV3-02).
Phase 140 Plan 03 — bridge status badge render-parity tests (BRIDGE-03).

Presence-based tests (per feedback_report_render_tests_presence_not_appearance):
assert the new SNMP/Bridge Status column header and label strings are PRESENT
in HTML and DOCX output, not that the visual layout matches. Mirrors the
project's established render-parity convention.

Node IDs:
  test_html_hardware_table_has_snmp_column
  test_html_hardware_table_snmp_failed_label
  test_docx_hardware_table_has_snmp_column
  test_docx_hardware_table_snmp_failed_label
  test_html_hardware_table_snmp_absent_renders_em_dash
  test_html_bridge_status_labels_and_caveat_present
  test_docx_bridge_status_labels_and_caveat_present
  test_bridge_status_content_model_parity_html_docx
"""
from __future__ import annotations

from types import SimpleNamespace


def _make_minimal_cfg():
    """Minimal cfg SimpleNamespace mirroring test_docx_report.py pattern."""
    return SimpleNamespace(
        assessment=SimpleNamespace(
            name="Test Org",
            report_owner="Test Owner",
            data_classification="CONFIDENTIAL",
            timezone="UTC",
            logo_path=None,
        ),
        output=SimpleNamespace(directory="/tmp/quirk_test_report_parity"),
    )


def _make_exec_content(hardware_devices):
    from quirk.reports.content_model import ExecContent

    return ExecContent(
        narrative_lead="Test narrative lead.",
        narrative_drivers=[],
        top_risks=[],
        roadmap_items=[],
        score_total=80,
        score_band="GOOD",
        subscores={},
        raw_sum=0,
        sev_counts={},
        hardware_devices=hardware_devices,
    )


_HW_AUTH_PRIV = {
    "remediation_tier": "Tier 2",
    "vendor": "Cisco",
    "model": "Catalyst 9300",
    "host": "10.0.0.1",
    "port": 161,
    "pqc_status": "unknown",
    "confidence": "high",
    "eol_date": None,
    "snmp_version": "v3 auth+priv",
}

_HW_FAILED_FALLBACK = {
    "remediation_tier": "Tier 3",
    "vendor": "Juniper",
    "model": "EX4300",
    "host": "10.0.0.2",
    "port": 161,
    "pqc_status": "unknown",
    "confidence": "medium",
    "eol_date": None,
    "snmp_version": "v3-failed-fell-back",
}

_HW_NO_SNMP_ATTEMPTED = {
    "remediation_tier": "Tier N/A",
    "vendor": "HP",
    "model": "ProCurve",
    "host": "10.0.0.3",
    "port": 22,
    "pqc_status": "unknown",
    "confidence": "low",
    "eol_date": None,
    # snmp_version absent entirely — SNMP never attempted for this device
}

_HW_BRIDGE_PARTIAL = {
    "remediation_tier": "Tier 2",
    "vendor": "Cisco",
    "model": "Catalyst 9300",
    "host": "10.0.1.1",
    "port": 161,
    "pqc_status": "unknown",
    "confidence": "high",
    "eol_date": None,
    "snmp_version": "v3 auth+priv",
    "bridge_status": "partial_only",
}

_HW_BRIDGE_CONFIRMED = {
    "remediation_tier": "Tier 2",
    "vendor": "Cisco",
    "model": "Legacy Switch",
    "host": "10.0.1.2",
    "port": 161,
    "pqc_status": "unknown",
    "confidence": "high",
    "eol_date": None,
    "snmp_version": "v3 auth+priv",
    "bridge_status": "upstream_mitigated",
}

_BRIDGE_CAVEAT_TEXT = (
    "Based on SNMP-derived network-path evidence; not independently confirmed"
    " by traffic inspection."
)


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


def test_html_hardware_table_has_snmp_column():
    """HTML hardware table renders an SNMP column header and the auth+priv label."""
    from quirk.reports.html_renderer import render_hardware_section

    html = render_hardware_section([_HW_AUTH_PRIV])
    assert ">SNMP<" in html, "SNMP column header not present in HTML output"
    assert "v3 auth+priv" in html, "v3 auth+priv label not present in HTML output"


def test_html_hardware_table_snmp_failed_label():
    """A v3-failed-fell-back device renders 'v3 failed -> v2c', not plain 'v2c'."""
    from quirk.reports.html_renderer import render_hardware_section

    html = render_hardware_section([_HW_FAILED_FALLBACK])
    assert "v3 failed" in html, "v3 failed → v2c label not present in HTML output"
    assert "v3 failed → v2c" in html


def test_html_hardware_table_snmp_absent_renders_em_dash():
    """A device with no snmp_version renders the em dash, distinct from 'No SNMP'."""
    from quirk.reports.html_renderer import render_hardware_section

    html = render_hardware_section([_HW_NO_SNMP_ATTEMPTED])
    assert "<td>—</td>" in html, "em dash not rendered for absent snmp_version"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_docx_hardware_table_has_snmp_column(tmp_path):
    """DOCX hardware table renders an SNMP column header and the auth+priv label."""
    from docx import Document

    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "report.docx")
    exec_content = _make_exec_content([_HW_AUTH_PRIV])
    result = render_docx_report(
        path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content
    )
    assert result is True

    doc = Document(path)
    full_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "SNMP" in full_text, "SNMP column header not present in DOCX output"
    assert "v3 auth+priv" in full_text, "v3 auth+priv label not present in DOCX output"


def test_docx_hardware_table_snmp_failed_label(tmp_path):
    """A v3-failed-fell-back device renders 'v3 failed -> v2c' in DOCX output."""
    from docx import Document

    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "report.docx")
    exec_content = _make_exec_content([_HW_FAILED_FALLBACK])
    result = render_docx_report(
        path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content
    )
    assert result is True

    doc = Document(path)
    full_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "v3 failed → v2c" in full_text


# ---------------------------------------------------------------------------
# Bridge status badge (Phase 140 BRIDGE-03)
# ---------------------------------------------------------------------------


def test_html_bridge_status_labels_and_caveat_present():
    """HTML renders both bridge badge labels and the verbatim Pitfall-3 caveat."""
    from quirk.reports.html_renderer import render_hardware_section

    html = render_hardware_section([_HW_BRIDGE_PARTIAL, _HW_BRIDGE_CONFIRMED])
    assert ">Bridge Status<" in html, "Bridge Status column header not present in HTML output"
    assert "Partial (assumed)" in html, "partial_only label not present in HTML output"
    assert "SNMP-confirmed" in html, "upstream_mitigated label not present in HTML output"
    assert "partial_only" not in html, "raw enum string leaked into HTML output"
    assert "upstream_mitigated" not in html, "raw enum string leaked into HTML output"
    assert _BRIDGE_CAVEAT_TEXT in html, "verbatim Pitfall-3 caveat not present in HTML output"


def test_docx_bridge_status_labels_and_caveat_present(tmp_path):
    """DOCX renders both bridge badge labels and the caveat as always-visible body text."""
    from docx import Document

    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "report.docx")
    exec_content = _make_exec_content([_HW_BRIDGE_PARTIAL, _HW_BRIDGE_CONFIRMED])
    result = render_docx_report(
        path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content
    )
    assert result is True

    doc = Document(path)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Bridge Status" in table_text, "Bridge Status column header not present in DOCX output"
    assert "Partial (assumed)" in table_text, "partial_only label not present in DOCX output"
    assert "SNMP-confirmed" in table_text, "upstream_mitigated label not present in DOCX output"
    # Caveat must be always-visible body text (a real paragraph), not hidden/collapsed.
    assert _BRIDGE_CAVEAT_TEXT in body_text, (
        "verbatim Pitfall-3 caveat not present as always-visible DOCX body text"
    )


def test_bridge_status_content_model_parity_html_docx(tmp_path):
    """HTML and DOCX surface the same bridge labels + caveat (content-model parity)."""
    from docx import Document

    from quirk.reports.docx_renderer import render_docx_report
    from quirk.reports.html_renderer import render_hardware_section

    html = render_hardware_section([_HW_BRIDGE_PARTIAL, _HW_BRIDGE_CONFIRMED])

    path = str(tmp_path / "report.docx")
    exec_content = _make_exec_content([_HW_BRIDGE_PARTIAL, _HW_BRIDGE_CONFIRMED])
    result = render_docx_report(
        path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content
    )
    assert result is True
    doc = Document(path)
    docx_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )

    for label in ("Partial (assumed)", "SNMP-confirmed"):
        assert label in html, f"{label!r} missing from HTML"
        assert label in docx_text, f"{label!r} missing from DOCX"
    assert _BRIDGE_CAVEAT_TEXT in html, "caveat missing from HTML"
    assert _BRIDGE_CAVEAT_TEXT in docx_text, "caveat missing from DOCX"
