"""Phase 100 Plan 02 — DOCX renderer tests (FMT-03 / D-09..D-12).

Tests structural fidelity, graceful-skip, and cross-surface content parity
for quirk/reports/docx_renderer.py.

Node IDs:
  test_docx_written
  test_logo_placeholder
  test_docx_section_headings
  test_docx_findings_table_cols
  test_docx_graceful_skip
  test_docx_skip_advisory
"""
from __future__ import annotations

import os
import sys
from types import SimpleNamespace

import pytest


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def _make_minimal_cfg():
    """Minimal cfg SimpleNamespace mirroring test_html_report.py pattern."""
    return SimpleNamespace(
        assessment=SimpleNamespace(
            name="Test Org",
            report_owner="Test Owner",
            data_classification="CONFIDENTIAL",
            timezone="UTC",
            logo_path=None,
        ),
        output=SimpleNamespace(directory="/tmp/quirk_test_docx"),
    )


# ---------------------------------------------------------------------------
# Tests: DOCX written and structurally valid
# ---------------------------------------------------------------------------

def test_docx_written(tmp_path):
    """render_docx_report writes a valid .docx file and returns True."""
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "report.docx")
    result = render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[])
    assert result is True, "render_docx_report returned False — expected True"
    assert os.path.exists(path), f"DOCX file not created at {path}"
    # Must be openable as a valid Word document
    doc = Document(path)
    assert len(doc.paragraphs) > 0, "DOCX has no paragraphs"


def test_logo_placeholder(tmp_path):
    """First paragraph of the DOCX is the verbatim logo placeholder string."""
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "logo_test.docx")
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[])
    doc = Document(path)
    # First paragraph must be the exact locked string (D-12 / 100-UI-SPEC.md §C)
    assert doc.paragraphs[0].text == "[ Insert organization logo here ]", (
        f"First paragraph is {doc.paragraphs[0].text!r}, "
        f"expected '[ Insert organization logo here ]'"
    )


def test_docx_section_headings(tmp_path):
    """DOCX contains Heading-1 paragraphs for all four required sections."""
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "headings_test.docx")
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[])
    doc = Document(path)

    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    required = {"Executive Summary", "Findings", "Remediation Roadmap", "Score Breakdown"}
    for h in required:
        assert h in heading_texts, (
            f"Heading-1 '{h}' not found in DOCX. Found headings: {heading_texts}"
        )


def test_docx_findings_table_cols(tmp_path):
    """First Word table (Findings) has 7 columns with the locked header order."""
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "table_test.docx")
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[])
    doc = Document(path)

    # Find the findings table — it's the first 7-column table
    findings_table = None
    for tbl in doc.tables:
        if len(tbl.columns) == 7:
            findings_table = tbl
            break
    assert findings_table is not None, "No 7-column findings table found in DOCX"

    expected_headers = [
        "Severity", "Title", "Host", "Port",
        "Description", "Recommendation", "Quantum Risk",
    ]
    actual_headers = [cell.text for cell in findings_table.rows[0].cells]
    assert actual_headers == expected_headers, (
        f"Findings table headers mismatch.\nExpected: {expected_headers}\nGot: {actual_headers}"
    )


def test_docx_empty_findings_row(tmp_path):
    """When findings is empty, the findings table has header + 'No findings recorded' row."""
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "empty_findings_test.docx")
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[])
    doc = Document(path)

    findings_table = None
    for tbl in doc.tables:
        if len(tbl.columns) == 7:
            findings_table = tbl
            break
    assert findings_table is not None, "No 7-column findings table found in DOCX"

    # Should have header row + 1 data row with the empty-state message
    assert len(findings_table.rows) == 2, (
        f"Expected 2 rows (header + empty-state), got {len(findings_table.rows)}"
    )
    first_data_cell = findings_table.rows[1].cells[0].text
    assert first_data_cell == "No findings recorded for this scan.", (
        f"Empty-state row first cell is {first_data_cell!r}, "
        f"expected 'No findings recorded for this scan.'"
    )


# ---------------------------------------------------------------------------
# Tests: graceful skip when python-docx is absent
# ---------------------------------------------------------------------------

def test_docx_graceful_skip(monkeypatch, tmp_path, capsys):
    """Returns False and does not raise when docx module is absent in sys.modules."""
    monkeypatch.setitem(sys.modules, "docx", None)
    # Re-import to exercise the lazy import path
    if "quirk.reports.docx_renderer" in sys.modules:
        del sys.modules["quirk.reports.docx_renderer"]
    from quirk.reports.docx_renderer import render_docx_report

    result = render_docx_report(
        path=str(tmp_path / "should_not_exist.docx"),
        cfg=_make_minimal_cfg(),
        findings=[],
    )
    assert result is False, f"Expected False when docx absent, got {result!r}"


def test_docx_skip_advisory(monkeypatch, tmp_path, capsys):
    """Prints the verbatim advisory to stderr when python-docx is absent."""
    monkeypatch.setitem(sys.modules, "docx", None)
    if "quirk.reports.docx_renderer" in sys.modules:
        del sys.modules["quirk.reports.docx_renderer"]
    from quirk.reports.docx_renderer import render_docx_report

    render_docx_report(
        path=str(tmp_path / "advisory_test.docx"),
        cfg=_make_minimal_cfg(),
        findings=[],
    )
    captured = capsys.readouterr()
    assert "python-docx is not installed" in captured.err, (
        f"Advisory not in stderr. stderr was: {captured.err!r}"
    )
    assert "pip install quirk-scanner[docx]" in captured.err, (
        f"Install hint not in stderr. stderr was: {captured.err!r}"
    )
