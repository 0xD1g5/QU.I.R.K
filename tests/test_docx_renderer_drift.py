"""Phase 156 (HWLC-10/HWLC-11 / D-11 / D-13 / T-156-04 / T-156-13) — DOCX
"Recent Lifecycle Changes" drift section tests.

Mirrors tests/test_docx_report.py's Document-open + text-readback idiom.
"""
from __future__ import annotations

from types import SimpleNamespace


def _make_minimal_cfg():
    return SimpleNamespace(
        assessment=SimpleNamespace(
            name="Drift Test Org",
            report_owner="Drift Tester",
            data_classification="CONFIDENTIAL",
            timezone="UTC",
            logo_path=None,
        ),
        output=SimpleNamespace(directory="/tmp/quirk_test_docx_drift"),
    )


def _base_exec_content(hardware_drift_events, hardware_devices=None):
    from quirk.reports.content_model import ExecContent

    return ExecContent(
        narrative_lead="Test narrative lead.",
        narrative_drivers=[],
        top_risks=[],
        roadmap_items=[],
        score_total=70,
        score_band="FAIR",
        subscores={},
        raw_sum=0,
        sev_counts={},
        hardware_drift_events=hardware_drift_events,
        hardware_devices=hardware_devices or [],
    )


def _base_device(**overrides) -> dict:
    device = {
        "host": "10.0.0.5",
        "port": 22,
        "vendor": "Schneider Electric",
        "model": "M221",
        "pqc_status": "unsupported",
        "confidence": "high",
        "eol_date": None,
        "remediation_tier": "Tier 1",
    }
    device.update(overrides)
    return device


def _base_event(**overrides) -> dict:
    event = {
        "host": "10.20.30.5",
        "port": 502,
        "event_type": "tier_crossing",
        "old_value": "Tier 2",
        "new_value": "Tier 1",
        "direction": "worsened",
        "detected_at": "2026-08-14",
        "vendor": "Schneider Electric",
        "model": "M221",
    }
    event.update(overrides)
    return event


def _all_paragraph_texts(doc):
    return [p.text for p in doc.paragraphs]


def test_docx_drift_section_advisory_caption_unconditional(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report, _DRIFT_ADVISORY_CAPTION

    path = str(tmp_path / "drift.docx")
    exec_content = _base_exec_content([_base_event()])
    result = render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)
    assert result is True

    doc = Document(path)
    texts = _all_paragraph_texts(doc)
    assert _DRIFT_ADVISORY_CAPTION in texts
    assert "Advisory — hardware lifecycle changes do not affect the readiness score." in texts


def test_docx_drift_section_heading_present(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "drift_heading.docx")
    exec_content = _base_exec_content([_base_event()])
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

    doc = Document(path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Recent Lifecycle Changes" in headings


def test_docx_drift_table_has_header_plus_one_row_per_event(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "drift_table.docx")
    events = [_base_event(), _base_event(host="10.20.30.6", event_type="cve_delta", direction="neutral")]
    exec_content = _base_exec_content(events)
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

    doc = Document(path)
    # Find the drift table by its header row.
    drift_tables = [
        t for t in doc.tables
        if t.rows[0].cells[0].text == "Device"
    ]
    assert len(drift_tables) == 1
    tbl = drift_tables[0]
    assert len(tbl.rows) == 1 + len(events)
    header_texts = [c.text for c in tbl.rows[0].cells]
    assert header_texts == ["Device", "Change", "Transition", "Direction", "Detected"]


def test_docx_drift_table_escapes_script_as_literal_text(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "drift_script.docx")
    event = _base_event(old_value="<script>alert(1)</script>", new_value="Tier 1")
    exec_content = _base_exec_content([event])
    result = render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)
    assert result is True

    doc = Document(path)
    drift_tables = [t for t in doc.tables if t.rows[0].cells[0].text == "Device"]
    assert len(drift_tables) == 1
    transition_cell_text = drift_tables[0].rows[1].cells[2].text
    assert "<script>alert(1)</script>" in transition_cell_text


def test_docx_no_drift_section_when_events_empty(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "no_drift.docx")
    exec_content = _base_exec_content([])
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

    doc = Document(path)
    texts = _all_paragraph_texts(doc)
    assert not any("Recent Lifecycle Changes" in t for t in texts)
    assert not any(t.rows[0].cells[0].text == "Device" for t in doc.tables)


# ---------------------------------------------------------------------------
# Phase 157 (HWLC-18) — EOL/Tier forecast subsection
# ---------------------------------------------------------------------------


def _base_exec_content_with_forecast(eol_forecast, hardware_drift_events=None):
    from quirk.reports.content_model import ExecContent

    return ExecContent(
        narrative_lead="Test narrative lead.",
        narrative_drivers=[],
        top_risks=[],
        roadmap_items=[],
        score_total=70,
        score_band="FAIR",
        subscores={},
        raw_sum=0,
        sev_counts={},
        hardware_drift_events=hardware_drift_events or [],
        eol_forecast=eol_forecast,
    )


def _two_bucket_forecast(**overrides):
    forecast = {
        "narrative": "Projected EOL/tier posture over the next 12 months.",
        "buckets": [
            {
                "label": "0-6 months",
                "count": 2,
                "tier_counts": {"Tier 1": 2},
                "sentence": "2 device(s) reach EOL or a worse CNSA 2.0 tier within 0-6 months.",
            },
            {
                "label": "7-12 months",
                "count": 1,
                "tier_counts": {"Tier 2": 1},
                "sentence": "1 device(s) reach EOL or a worse CNSA 2.0 tier within 7-12 months.",
            },
        ],
        "catalog_last_verified": "2026-05-01",
        "catalog_stale": False,
        "total_devices_with_eol": 3,
    }
    forecast.update(overrides)
    return forecast


def test_docx_forecast_subsection_present(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "forecast.docx")
    exec_content = _base_exec_content_with_forecast(_two_bucket_forecast())
    result = render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)
    assert result is True

    doc = Document(path)
    texts = _all_paragraph_texts(doc)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "EOL/Tier Forecast" in headings
    assert "2 device(s) reach EOL or a worse CNSA 2.0 tier within 0-6 months." in texts
    assert "1 device(s) reach EOL or a worse CNSA 2.0 tier within 7-12 months." in texts


def test_docx_forecast_absent_when_empty(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    for empty_forecast in ({}, {"narrative": "", "buckets": []}):
        path = str(tmp_path / f"forecast_empty_{id(empty_forecast)}.docx")
        exec_content = _base_exec_content_with_forecast(empty_forecast)
        render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

        doc = Document(path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "EOL/Tier Forecast" not in headings


def test_docx_forecast_renders_with_zero_drift_events(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "forecast_zero_drift.docx")
    exec_content = _base_exec_content_with_forecast(_two_bucket_forecast(), hardware_drift_events=[])
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

    doc = Document(path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "EOL/Tier Forecast" in headings
    assert "Recent Lifecycle Changes" not in headings


def test_docx_forecast_uses_paragraphs_not_table(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path_empty = str(tmp_path / "forecast_baseline.docx")
    exec_content_empty = _base_exec_content_with_forecast({})
    render_docx_report(path=path_empty, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content_empty)
    doc_empty = Document(path_empty)
    table_count_empty = len(doc_empty.tables)

    path_forecast = str(tmp_path / "forecast_populated.docx")
    exec_content = _base_exec_content_with_forecast(_two_bucket_forecast())
    render_docx_report(path=path_forecast, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)
    doc_forecast = Document(path_forecast)

    assert len(doc_forecast.tables) == table_count_empty


def test_docx_forecast_surfaces_stale_catalog(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report

    path_stale = str(tmp_path / "forecast_stale.docx")
    stale_forecast = _two_bucket_forecast(catalog_stale=True, catalog_last_verified="2025-01-01")
    exec_content_stale = _base_exec_content_with_forecast(stale_forecast)
    render_docx_report(path=path_stale, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content_stale)
    doc_stale = Document(path_stale)
    texts_stale = _all_paragraph_texts(doc_stale)
    assert any("2025-01-01" in t for t in texts_stale)

    path_fresh = str(tmp_path / "forecast_fresh.docx")
    fresh_forecast = _two_bucket_forecast(catalog_stale=False, catalog_last_verified="2025-01-01")
    exec_content_fresh = _base_exec_content_with_forecast(fresh_forecast)
    render_docx_report(path=path_fresh, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content_fresh)
    doc_fresh = Document(path_fresh)
    texts_fresh = _all_paragraph_texts(doc_fresh)
    assert not any("2025-01-01" in t for t in texts_fresh)


# ---------------------------------------------------------------------------
# Phase 159 HWLC-13/D-159-M/N/P: always-visible "partial re-probe" banner.
# Import the banner constant rather than retyping the literal, so a future
# copy change cannot silently pass a stale assertion.
# ---------------------------------------------------------------------------


def test_docx_drift_section_checkin_shows_banner(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report, _PARTIAL_SCAN_BANNER

    path = str(tmp_path / "drift_checkin.docx")
    exec_content = _base_exec_content([_base_event(is_partial_scan=True)])
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

    doc = Document(path)
    texts = _all_paragraph_texts(doc)
    assert _PARTIAL_SCAN_BANNER in texts


def test_docx_drift_section_no_checkin_omits_banner(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report, _PARTIAL_SCAN_BANNER

    path = str(tmp_path / "drift_no_checkin.docx")
    exec_content = _base_exec_content([_base_event(is_partial_scan=False)])
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

    doc = Document(path)
    texts = _all_paragraph_texts(doc)
    assert _PARTIAL_SCAN_BANNER not in texts


def test_docx_hardware_section_checkin_shows_banner(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report, _PARTIAL_SCAN_BANNER

    path = str(tmp_path / "hardware_checkin.docx")
    exec_content = _base_exec_content(
        [], hardware_devices=[_base_device(is_partial_scan=True)]
    )
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

    doc = Document(path)
    texts = _all_paragraph_texts(doc)
    assert _PARTIAL_SCAN_BANNER in texts


def test_docx_hardware_section_no_checkin_omits_banner(tmp_path):
    from docx import Document
    from quirk.reports.docx_renderer import render_docx_report, _PARTIAL_SCAN_BANNER

    path = str(tmp_path / "hardware_no_checkin.docx")
    exec_content = _base_exec_content(
        [], hardware_devices=[_base_device(is_partial_scan=False)]
    )
    render_docx_report(path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content)

    doc = Document(path)
    texts = _all_paragraph_texts(doc)
    assert _PARTIAL_SCAN_BANNER not in texts
