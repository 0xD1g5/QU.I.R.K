"""Phase 141 Plan 06 — Modbus/BACnet column render-parity tests (OTICS-05).

Presence-based tests (per feedback_report_render_tests_presence_not_appearance):
assert the new Modbus/BACnet column headers, label strings, and the D-13
abort caveat are PRESENT in HTML and DOCX output, not that the visual layout
matches. Mirrors the project's established render-parity convention
(test_report_render_parity.py's SNMP/Bridge Status precedent).

Node IDs:
  test_html_hardware_table_has_modbus_and_bacnet_columns
  test_html_hardware_table_otics_abort_caveat_present
  test_html_hardware_table_otics_not_attempted_em_dash
  test_docx_hardware_table_has_modbus_and_bacnet_columns
  test_docx_hardware_table_otics_abort_caveat_present
  test_otics_content_model_parity_html_docx
"""
from __future__ import annotations

from types import SimpleNamespace


def _make_minimal_cfg():
    """Minimal cfg SimpleNamespace mirroring test_report_render_parity.py pattern."""
    return SimpleNamespace(
        assessment=SimpleNamespace(
            name="Test Org",
            report_owner="Test Owner",
            data_classification="CONFIDENTIAL",
            timezone="UTC",
            logo_path=None,
        ),
        output=SimpleNamespace(directory="/tmp/quirk_test_report_otics_columns"),
    )


def _make_exec_content(hardware_devices):
    from quirk.reports.content_model import ExecContent

    return ExecContent(
        narrative_lead="Test narrative lead.",
        narrative_drivers=[],
        top_risks=[],
        roadmap_items=[],
        score_total=80,
        score_band="GOOD",
        subscores={},
        raw_sum=0,
        sev_counts={},
        hardware_devices=hardware_devices,
    )


_HW_MODBUS_IDENTIFIED = {
    "remediation_tier": "Tier 2",
    "vendor": "Schneider Electric",
    "model": "M221",
    "host": "10.0.2.1",
    "port": 502,
    "pqc_status": "unknown",
    "confidence": "high",
    "eol_date": None,
    "modbus_vendor": "Schneider Electric",
    "modbus_model": "M221",
    "modbus_probe_state": "identified",
    "bacnet_probe_state": None,
}

_HW_BACNET_IDENTIFIED = {
    "remediation_tier": "Tier 2",
    "vendor": "Johnson Controls",
    "model": "FX-PCG",
    "host": "10.0.2.2",
    "port": 47808,
    "pqc_status": "unknown",
    "confidence": "high",
    "eol_date": None,
    "modbus_probe_state": None,
    "bacnet_vendor": "Johnson Controls",
    "bacnet_model": "FX-PCG",
    "bacnet_probe_state": "identified",
}

_HW_MODBUS_ABORTED = {
    "remediation_tier": "Tier 3",
    "vendor": "Unknown",
    "model": None,
    "host": "10.0.2.3",
    "port": 502,
    "pqc_status": "unknown",
    "confidence": "low",
    "eol_date": None,
    "modbus_probe_state": "aborted_anomalous_response",
    "bacnet_probe_state": None,
}

_HW_NOT_ATTEMPTED = {
    "remediation_tier": "Tier N/A",
    "vendor": "HP",
    "model": "ProCurve",
    "host": "10.0.2.4",
    "port": 22,
    "pqc_status": "unknown",
    "confidence": "low",
    "eol_date": None,
    # modbus_probe_state / bacnet_probe_state absent entirely — never attempted
}

_HW_NO_RESPONSE = {
    "remediation_tier": "Tier 3",
    "vendor": "Unknown",
    "model": None,
    "host": "10.0.2.5",
    "port": 502,
    "pqc_status": "unknown",
    "confidence": "low",
    "eol_date": None,
    "modbus_probe_state": "no_response",
    "bacnet_probe_state": "no_match",
}

_OTICS_ABORT_CAVEAT_TEXT = (
    "Modbus/BACnet probe aborted — anomalous response. The device returned a"
    " malformed frame, reset the connection, or timed out; QU.I.R.K. stopped"
    " probing this host per its one-strike safety policy. Worth a closer"
    " manual look."
)


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


def test_html_hardware_table_has_modbus_and_bacnet_columns():
    """HTML hardware table renders Modbus/BACnet column headers and identified labels."""
    from quirk.reports.html_renderer import render_hardware_section

    html = render_hardware_section([_HW_MODBUS_IDENTIFIED, _HW_BACNET_IDENTIFIED])
    assert ">Modbus<" in html, "Modbus column header not present in HTML output"
    assert ">BACnet<" in html, "BACnet column header not present in HTML output"
    assert "<td>Modbus</td>" in html, "Modbus identified label not present in HTML output"
    assert "<td>BACnet</td>" in html, "BACnet identified label not present in HTML output"


def test_html_hardware_table_otics_abort_caveat_present():
    """A device with an aborted Modbus/BACnet probe renders 'Probe aborted' + the caveat."""
    from quirk.reports.html_renderer import render_hardware_section

    html = render_hardware_section([_HW_MODBUS_ABORTED])
    assert "Probe aborted" in html, "Probe aborted label not present in HTML output"
    assert _OTICS_ABORT_CAVEAT_TEXT in html, "abort caveat sentence not present in HTML output"


def test_html_hardware_table_otics_not_attempted_em_dash():
    """A device with no modbus/bacnet_probe_state renders the em dash for both columns,
    distinct from 'No response'/'No match'."""
    from quirk.reports.html_renderer import render_hardware_section

    html = render_hardware_section([_HW_NOT_ATTEMPTED])
    assert "<td>—</td>" in html, "em dash not rendered for absent probe_state"
    assert "No response" not in html
    assert "No match" not in html

    html_no_response = render_hardware_section([_HW_NO_RESPONSE])
    assert "No response" in html_no_response
    assert "No match" in html_no_response


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_docx_hardware_table_has_modbus_and_bacnet_columns(tmp_path):
    """DOCX hardware table renders Modbus/BACnet column headers and identified labels."""
    from docx import Document

    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "report.docx")
    exec_content = _make_exec_content([_HW_MODBUS_IDENTIFIED, _HW_BACNET_IDENTIFIED])
    result = render_docx_report(
        path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content
    )
    assert result is True

    doc = Document(path)
    full_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Modbus" in full_text, "Modbus column header/label not present in DOCX output"
    assert "BACnet" in full_text, "BACnet column header/label not present in DOCX output"


def test_docx_hardware_table_otics_abort_caveat_present(tmp_path):
    """DOCX renders 'Probe aborted' + the caveat as always-visible body text."""
    from docx import Document

    from quirk.reports.docx_renderer import render_docx_report

    path = str(tmp_path / "report.docx")
    exec_content = _make_exec_content([_HW_MODBUS_ABORTED])
    result = render_docx_report(
        path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content
    )
    assert result is True

    doc = Document(path)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Probe aborted" in table_text, "Probe aborted label not present in DOCX output"
    assert _OTICS_ABORT_CAVEAT_TEXT in body_text, (
        "abort caveat sentence not present as always-visible DOCX body text"
    )


def test_otics_content_model_parity_html_docx(tmp_path):
    """HTML and DOCX surface the same Modbus/BACnet labels + caveat (content-model parity)."""
    from docx import Document

    from quirk.reports.docx_renderer import render_docx_report
    from quirk.reports.html_renderer import render_hardware_section

    devices = [_HW_MODBUS_IDENTIFIED, _HW_BACNET_IDENTIFIED, _HW_MODBUS_ABORTED]
    html = render_hardware_section(devices)

    path = str(tmp_path / "report.docx")
    exec_content = _make_exec_content(devices)
    result = render_docx_report(
        path=path, cfg=_make_minimal_cfg(), findings=[], exec_content=exec_content
    )
    assert result is True
    doc = Document(path)
    docx_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )

    for label in ("Modbus", "BACnet", "Probe aborted"):
        assert label in html, f"{label!r} missing from HTML"
        assert label in docx_text, f"{label!r} missing from DOCX"
    assert _OTICS_ABORT_CAVEAT_TEXT in html, "caveat missing from HTML"
    assert _OTICS_ABORT_CAVEAT_TEXT in docx_text, "caveat missing from DOCX"
